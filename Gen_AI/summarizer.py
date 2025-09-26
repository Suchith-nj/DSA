import os
import openai
import PyPDF2
from docx import Document
import streamlit as st
from io import BytesIO
import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter

class PureLLMSummarizer:
    def __init__(self, api_key: str, model: str = "gpt-3.5-turbo"):
        self.client = openai.OpenAI(api_key=api_key)
        self.model = model
        # tiktoken library is used to retrieve the appropriate token encoding for a given OpenAI model.
        self.encoding = tiktoken.encoding_for_model(model)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=3000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_path_or_bytes):
        """Extract text from PDF file"""
        try:
            if isinstance(file_path_or_bytes, str):
                with open(file_path_or_bytes, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            else:
                reader = PyPDF2.PdfReader(file_path_or_bytes)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path_or_bytes):
        """Extract text from DOCX file"""
        try:
            if isinstance(file_path_or_bytes, str):
                doc = Document(file_path_or_bytes)
            else:
                doc = Document(file_path_or_bytes)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path_or_bytes):
        """Extract text from TXT file"""
        try:
            if isinstance(file_path_or_bytes, str):
                with open(file_path_or_bytes, 'r', encoding='utf-8') as file:
                    return file.read()
            else:
                return file_path_or_bytes.read().decode('utf-8')
        except Exception as e:
            raise Exception(f"Error reading TXT: {str(e)}")
    
    def extract_text(self, file_path_or_bytes, file_type):
        """Extract text based on file type"""
        if file_type.lower() == 'pdf':
            return self.extract_text_from_pdf(file_path_or_bytes)
        elif file_type.lower() in ['docx', 'doc']:
            return self.extract_text_from_docx(file_path_or_bytes)
        elif file_type.lower() == 'txt':
            return self.extract_text_from_txt(file_path_or_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def count_tokens(self, text):
        """Count tokens in text"""
        return len(self.encoding.encode(text))
    
    def chunk_text(self, text):
        """Use LangChain's smart text splitter"""
        return self.text_splitter.split_text(text)
    
    def summarize_text(self, text, custom_prompt=None, strategy="auto"):
        """Summarize text using OpenAI"""
        
        # Default prompt
        if custom_prompt is None:
            custom_prompt = "Provide a comprehensive summary of the following text, highlighting the main points, key insights, and important conclusions:"
        
        token_count = self.count_tokens(text)
        
        # Choose strategy based on token count
        if strategy == "auto":
            if token_count < 3000:
                strategy = "single"
            else:
                strategy = "chunk"
        
        if strategy == "single":
            return self._summarize_single(text, custom_prompt)
        else:
            return self._summarize_chunks(text, custom_prompt)
    
    def _summarize_single(self, text, prompt):
        """Summarize text in one go"""
        messages = [
            {"role": "system", "content": "You are a helpful assistant that creates clear, concise summaries."},
            {"role": "user", "content": f"{prompt}\n\nText to summarize:\n{text}"}
        ]
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=messages,
            temperature=0.4,
            max_tokens=1000
        )
        
        return response.choices[0].message.content
    
    def _summarize_chunks(self, text, prompt):
        """Summarize large text by chunking"""
        chunks = self.chunk_text(text)
        
        # Step 1: Summarize each chunk
        chunk_summaries = []
        for i, chunk in enumerate(chunks):
            messages = [
                {"role": "system", "content": "You are a helpful assistant that creates clear, concise summaries."},
                {"role": "user", "content": f"Summarize the key points from this section (part {i+1} of {len(chunks)}):\n\n{chunk}"}
            ]
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.3,
                max_tokens=500
            )
            
            chunk_summaries.append(response.choices[0].message.content)
        
        # Step 2: Combine chunk summaries into final summary
        combined_summaries = "\n\n".join(chunk_summaries)
        
        final_messages = [
            {"role": "system", "content": "You are a helpful assistant that creates comprehensive summaries."},
            {"role": "user", "content": f"{prompt}\n\nPlease create a cohesive summary from these section summaries:\n\n{combined_summaries}"}
        ]
        
        final_response = self.client.chat.completions.create(
            model=self.model,
            messages=final_messages,
            temperature=0.3,
            max_tokens=1000
        )
        
        return final_response.choices[0].message.content
    
    def summarize_file(self, file_path_or_bytes, file_type, custom_prompt=None, strategy="auto"):
        """Complete pipeline: extract text and summarize"""
        # Extract text
        text = self.extract_text(file_path_or_bytes, file_type)
        
        if not text.strip():
            raise ValueError("No text found in the document")
        
        # Summarize
        summary = self.summarize_text(text, custom_prompt, strategy)
        return summary

# Streamlit Interface
def main():
    st.title("Document Summarizer")
    st.write("Upload PDF, DOCX, or TXT files for AI-powered summaries")
    
    # Configuration
    with st.sidebar:
        st.header("Configuration")
        api_key = st.text_input("OpenAI API Key", type="password")
        
        model = st.selectbox(
            "Model",
            ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo"],
        )
        
        strategy = st.selectbox(
            "Strategy",
            ["auto", "single", "chunk"],
            help={
                "auto": "Automatically choose based on document size",
                "single": "Process entire document at once (faster, limited size)",
                "chunk": "Split into chunks then combine (handles large docs)"
            }
        )
        
        use_custom_prompt = st.checkbox("Custom Prompt")
        custom_prompt = ""
        if use_custom_prompt:
            custom_prompt = st.text_area(
                "Custom Prompt",
                "Please provide a comprehensive summary focusing on the main arguments and conclusions:"
            )
    
    if not api_key:
        st.warning("Please enter your OpenAI API key in the sidebar.")
        return
    
    # Initialize summarizer
    try:
        summarizer = PureLLMSummarizer(api_key, model)
    except Exception as e:
        st.error(f"Error initializing: {str(e)}")
        return
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=['pdf', 'docx', 'txt']
    )
    
    if uploaded_file:
        st.subheader(f" {uploaded_file.name}")
        
        # Get file type
        file_type = uploaded_file.name.split('.')[-1]
        
        try:
            with st.spinner("Extracting text and generating summary..."):
                # Use BytesIO for uploaded files
                file_bytes = BytesIO(uploaded_file.read())
                
                summary = summarizer.summarize_file(
                    file_bytes,
                    file_type,
                    custom_prompt if use_custom_prompt else None,
                    strategy
                )
            
            st.success("Summary generated!")
            
            with st.expander("Summary", expanded=True):
                st.write(summary)
            
            st.download_button(
                " Download Summary",
                summary,
                file_name=f"summary_{uploaded_file.name.split('.')[0]}.txt",
                mime="text/plain"
            )
            
        except Exception as e:
            st.error(f"Error: {str(e)}")

if __name__ == "__main__":
    main()